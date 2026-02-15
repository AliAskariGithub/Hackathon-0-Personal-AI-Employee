#!/usr/bin/env python3
"""
Agent Factory Bronze Tier - File Triage Agent
Implements the SKILL.md workflow to process files from Needs_Action.
Supports: .md, .txt, .docx files
AI Provider: Groq API
"""

import os
import time
from datetime import datetime
from pathlib import Path
from typing import Optional

# Third-party imports
from groq import Groq
from docx import Document
from dotenv import load_dotenv
from colorama import Fore, Back, Style, init

# Initialize colorama for Windows compatibility
init(autoreset=True)


class FileTriageAgent:
    """AI Agent that processes files according to the SKILL.md manual."""

    def __init__(self, base_path=None):
        """Initialize the agent with vault paths and Groq API."""
        self.base_path = Path(base_path) if base_path else Path(__file__).parent
        self.needs_action_path = self.base_path / "Vault" / "Needs_Action"
        self.done_path = self.base_path / "Vault" / "Done"
        self.dashboard_path = self.base_path / "Vault" / "Dashboard.md"
        self.skill_path = self.base_path / "skills" / "file-triage" / "SKILL.md"

        # Load environment variables
        load_dotenv()

        # Initialize Groq client
        self.groq_api_key = os.getenv("GROQ_API_KEY")
        self.groq_model = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
        self.max_tokens = int(os.getenv("MAX_TOKENS", "2048"))
        self.temperature = float(os.getenv("TEMPERATURE", "0.7"))

        if not self.groq_api_key:
            print(f"{Fore.YELLOW}[!] WARNING: GROQ_API_KEY not found in .env file{Style.RESET_ALL}")
            print(f"{Fore.YELLOW}            Agent will run in simulation mode{Style.RESET_ALL}")
            self.groq_client = None
        else:
            self.groq_client = Groq(api_key=self.groq_api_key)
            print(f"{Fore.GREEN}[OK] Groq API initialized: {Fore.CYAN}{self.groq_model}{Style.RESET_ALL}")

        # Verify paths exist
        self._verify_paths()

    def _verify_paths(self):
        """Verify all required paths exist."""
        required_paths = [
            self.needs_action_path,
            self.done_path,
            self.dashboard_path,
            self.skill_path
        ]

        for path in required_paths:
            if not path.exists():
                raise FileNotFoundError(f"Required path does not exist: {path}")

    def monitor(self):
        """STEP 1: MONITOR - Identify files requiring processing."""
        print(f"\n{Fore.BLUE}[STEP 1: MONITOR]{Style.RESET_ALL}")

        # Get all supported files in Needs_Action
        supported_extensions = ['*.md', '*.txt', '*.docx']
        all_files = []

        for ext in supported_extensions:
            all_files.extend(self.needs_action_path.glob(ext))

        # Sort by modification time (FIFO)
        all_files = sorted(all_files, key=lambda p: p.stat().st_mtime)

        if not all_files:
            print(f"  {Fore.YELLOW}[EMPTY] No files found in Needs_Action/{Style.RESET_ALL}")
            return None

        selected_file = all_files[0]
        print(f"  {Fore.GREEN}[OK] Selected: {Fore.CYAN}{selected_file.name}{Fore.GREEN} ({selected_file.suffix}){Style.RESET_ALL}")
        return selected_file

    def read_file_content(self, file_path: Path) -> Optional[str]:
        """Read content from different file types."""
        try:
            if file_path.suffix.lower() == '.docx':
                # Read .docx file
                doc = Document(file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                return content
            elif file_path.suffix.lower() in ['.md', '.txt']:
                # Read text-based files
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                print(f"  {Fore.RED}[X] ERROR: Unsupported file type: {file_path.suffix}{Style.RESET_ALL}")
                return None
        except Exception as e:
            print(f"  {Fore.RED}[X] ERROR: Failed to read file: {e}{Style.RESET_ALL}")
            return None

    def process(self, file_path):
        """STEP 2: PROCESS - Read, analyze, and perform the task."""
        print(f"\n{Fore.BLUE}[STEP 2: PROCESS]{Style.RESET_ALL}")

        # Read the file content
        content = self.read_file_content(file_path)
        if not content:
            return None

        print(f"  {Fore.GREEN}[OK] Read {Fore.CYAN}{len(content)}{Fore.GREEN} characters from {Fore.CYAN}{file_path.name}{Style.RESET_ALL}")

        # Analyze the intent
        print(f"  {Fore.YELLOW}[ANALYZE] Analyzing request...{Style.RESET_ALL}")
        task_type = self._classify_task(content)
        print(f"  {Fore.GREEN}[OK] Task type: {Fore.MAGENTA}{task_type}{Style.RESET_ALL}")

        # Perform the task
        print(f"  {Fore.YELLOW}[AI] Generating response...{Style.RESET_ALL}")
        response = self._generate_response(content, task_type)

        if response:
            print(f"  {Fore.GREEN}[OK] Response generated successfully{Style.RESET_ALL}")
            return response
        else:
            print(f"  {Fore.RED}[X] ERROR: Failed to generate response{Style.RESET_ALL}")
            return None

    def _classify_task(self, content):
        """Classify the type of task based on content."""
        content_lower = content.lower()

        if any(word in content_lower for word in ['?', 'what', 'how', 'why', 'when', 'where']):
            return "Question"
        elif any(word in content_lower for word in ['analyze', 'review', 'examine']):
            return "Analysis"
        elif any(word in content_lower for word in ['write', 'create', 'generate', 'draft']):
            return "Writing"
        elif any(word in content_lower for word in ['code', 'function', 'script', 'program']):
            return "Code"
        elif any(word in content_lower for word in ['calculate', 'compute', 'solve']):
            return "Calculation"
        else:
            return "General"

    def _generate_response(self, content, task_type):
        """Generate an AI response using Groq API."""
        # If Groq client is not available, use simulation mode
        if not self.groq_client:
            return self._generate_simulated_response(content, task_type)

        try:
            # Create the prompt for Groq
            system_prompt = """You are a professional AI assistant working in an automated file triage system.
Your role is to read user requests from files and provide helpful, accurate, and well-structured responses.

Guidelines:
- Be professional and clear
- Provide actionable information
- Use proper formatting (markdown)
- Be thorough but concise
- If you cannot fulfill a request, explain why"""

            user_prompt = f"""Task Type: {task_type}

User Request:
{content}

Please provide a comprehensive response to this request."""

            # Call Groq API
            chat_completion = self.groq_client.chat.completions.create(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                model=self.groq_model,
                temperature=self.temperature,
                max_tokens=self.max_tokens,
            )

            response = chat_completion.choices[0].message.content
            return response

        except Exception as e:
            print(f"  {Fore.RED}[X] ERROR: Groq API call failed: {e}{Style.RESET_ALL}")
            print(f"  {Fore.YELLOW}[!] Falling back to simulation mode{Style.RESET_ALL}")
            return self._generate_simulated_response(content, task_type)

    def _generate_simulated_response(self, content, task_type):
        """Generate a simulated response when Groq API is unavailable."""
        response = f"""This is a simulated AI response for a {task_type} task.

**Original Request Summary:**
The file contains a request that has been classified as a {task_type} task.

**Response:**
In a production implementation with Groq API configured, this section would contain:
- A detailed analysis of the request
- A comprehensive answer or solution
- Relevant examples, code, or explanations
- Any necessary warnings or limitations

**Configuration Note:**
To enable real AI processing:
1. Copy `.env.example` to `.env`
2. Add your Groq API key: `GROQ_API_KEY=your_key_here`
3. Optionally customize the model and parameters
4. Restart the agent

Get your free Groq API key at: https://console.groq.com/
"""
        return response

    def update(self, file_path, response, task_type):
        """STEP 3: UPDATE - Write the response back into the file."""
        print(f"\n{Fore.BLUE}[STEP 3: UPDATE]{Style.RESET_ALL}")

        try:
            # For .docx files, we need to append to the document
            if file_path.suffix.lower() == '.docx':
                return self._update_docx(file_path, response, task_type)
            else:
                # For text-based files (.md, .txt)
                return self._update_text_file(file_path, response, task_type)

        except Exception as e:
            print(f"  {Fore.RED}[X] ERROR: Failed to update file: {e}{Style.RESET_ALL}")
            return False

    def _update_text_file(self, file_path, response, task_type):
        """Update text-based files (.md, .txt)."""
        # Read current content
        with open(file_path, 'r', encoding='utf-8') as f:
            original_content = f.read()

        # Prepare the response section
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        response_section = f"""
---

## AI Response

**Processed:** {timestamp}

{response}

**Task Type:** {task_type}
**Status:** Completed
"""

        # Append response to original content
        updated_content = original_content + response_section

        # Write back to file
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(updated_content)

        print(f"  {Fore.GREEN}[OK] Updated {Fore.CYAN}{file_path.name}{Fore.GREEN} with AI response{Style.RESET_ALL}")
        return True

    def _update_docx(self, file_path, response, task_type):
        """Update .docx files by appending response."""
        doc = Document(file_path)

        # Add separator
        doc.add_paragraph("_" * 50)

        # Add response header
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        doc.add_heading("AI Response", level=2)
        doc.add_paragraph(f"Processed: {timestamp}")
        doc.add_paragraph()

        # Add response content
        for line in response.split('\n'):
            if line.strip():
                doc.add_paragraph(line)

        # Add metadata
        doc.add_paragraph()
        doc.add_paragraph(f"Task Type: {task_type}")
        doc.add_paragraph(f"Status: Completed")

        # Save the document
        doc.save(file_path)

        print(f"  {Fore.GREEN}[OK] Updated {Fore.CYAN}{file_path.name}{Fore.GREEN} with AI response{Style.RESET_ALL}")
        return True

    def finalize(self, file_path):
        """STEP 4: FINALIZE - Move to Done and update Dashboard."""
        print(f"\n{Fore.BLUE}[STEP 4: FINALIZE]{Style.RESET_ALL}")

        filename = file_path.name
        task_name = file_path.stem

        # Move file to Done
        try:
            destination = self.done_path / filename
            file_path.rename(destination)
            print(f"  {Fore.GREEN}[OK] Moved {Fore.CYAN}{filename}{Fore.GREEN} -> {Fore.MAGENTA}Done/{Style.RESET_ALL}")
        except Exception as e:
            print(f"  {Fore.RED}[X] ERROR: Failed to move file: {e}{Style.RESET_ALL}")
            return False

        # Update Dashboard
        try:
            self._update_dashboard(task_name, "Completed", filename)
            print(f"  {Fore.GREEN}[OK] Updated Dashboard: {Fore.CYAN}{task_name}{Fore.GREEN} -> {Fore.YELLOW}Completed{Style.RESET_ALL}")
        except Exception as e:
            print(f"  {Fore.RED}[X] ERROR: Failed to update Dashboard: {e}{Style.RESET_ALL}")
            return False

        # Log completion
        completion_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        print(f"\n{Fore.GREEN}[DONE] Task completed: {Fore.CYAN}{filename}{Fore.GREEN} at {Fore.WHITE}{completion_time}{Style.RESET_ALL}")

        return True

    def _update_dashboard(self, task_name, new_status, filename=None):
        """Update the status of a task in the Dashboard."""
        import re

        current_date = datetime.now().strftime("%Y-%m-%d")
        current_time = datetime.now().strftime("%H:%M:%S")
        current_timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Read dashboard
        with open(self.dashboard_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Determine status badge style and statistics update
        if new_status == "Completed":
            status_badge = '<span style="background: linear-gradient(135deg, #10b981 0%, #059669 100%); color: white; padding: 6px 16px; border-radius: 20px; font-size: 12px; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px; box-shadow: 0 2px 8px rgba(16, 185, 129, 0.4);">Completed</span>'
            stats_update = {'completed_delta': 1, 'pending_delta': -1}
        elif new_status == "Error":
            status_badge = '<span style="background: linear-gradient(135deg, #ef4444 0%, #dc2626 100%); color: white; padding: 6px 16px; border-radius: 20px; font-size: 12px; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px; box-shadow: 0 2px 8px rgba(239, 68, 68, 0.4);">Error</span>'
            stats_update = {'error_delta': 1, 'pending_delta': -1}
        else:
            status_badge = f'<span style="background: linear-gradient(135deg, #f59e0b 0%, #d97706 100%); color: white; padding: 6px 16px; border-radius: 20px; font-size: 12px; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px; box-shadow: 0 2px 8px rgba(245, 158, 11, 0.4);">{new_status}</span>'
            stats_update = {}

        # Find and update the task row
        # Pattern to match the entire row containing the task name
        pattern = rf'(<tr style="[^"]*">\s*<td style="padding: 16px; text-align: center; font-weight: 700; color: #6366f1; font-size: 16px;">\d+</td>\s*<td style="padding: 16px; color: #e4e6eb; font-size: 14px;">){re.escape(task_name)}(</td>\s*<td style="padding: 16px; text-align: center;">).*?(</td>\s*<td style="padding: 16px; text-align: center; color: #9ca3af; font-size: 14px;">).*?(</td>\s*<td style="padding: 16px; text-align: center; color: #9ca3af; font-size: 14px;">).*?(</td>\s*<td style="padding: 16px;">).*?(</td>\s*</tr>)'

        def replace_row(match):
            # Extract the filename from the existing link if not provided
            nonlocal filename
            if filename is None:
                link_match = re.search(r'href="[^"]*?/([^"]+)"', match.group(0))
                if link_match:
                    filename = link_match.group(1)
                else:
                    filename = f"{task_name}.md"  # fallback

            # Build the updated row
            return (
                f'{match.group(1)}{task_name}{match.group(2)}\n        {status_badge}\n      '
                f'{match.group(3)}{current_date}{match.group(4)}{current_time}'
                f'{match.group(5)}<a href="Done/{filename}" style="color: #6366f1; text-decoration: none; font-weight: 500; font-size: 14px;">📄 Done/{filename}</a>'
                f'{match.group(6)}'
            )

        content = re.sub(pattern, replace_row, content, flags=re.DOTALL)

        # Update statistics
        if stats_update:
            content = self._update_statistics(content, **stats_update)

        # Update last updated timestamp (if exists in footer)
        if 'Last Updated:' in content:
            content = re.sub(
                r'<p style="margin: 0; color: #[^"]+;"><strong>Last Updated:</strong> [^<]+</p>',
                f'<p style="margin: 0; color: #9ca3af;"><strong>Last Updated:</strong> {current_timestamp}</p>',
                content
            )

        # Write back
        with open(self.dashboard_path, 'w', encoding='utf-8') as f:
            f.write(content)

    def _update_statistics(self, content, total_delta=0, completed_delta=0, pending_delta=0, error_delta=0):
        """Update the statistics section in the Dashboard."""
        import re

        # Extract current statistics - updated pattern for new Dashboard format
        # Pattern: <h3 style="...font-size: 48px...">NUMBER</h3>
        total_match = re.search(
            r'<h3 style="[^"]*font-size: 48px[^"]*">(\d+)</h3>\s*<p style="[^"]*">Total Tasks</p>',
            content,
            re.DOTALL
        )
        completed_match = re.search(
            r'<h3 style="[^"]*font-size: 48px[^"]*">(\d+)</h3>\s*<p style="[^"]*">Completed</p>',
            content,
            re.DOTALL
        )
        pending_match = re.search(
            r'<h3 style="[^"]*font-size: 48px[^"]*">(\d+)</h3>\s*<p style="[^"]*">Pending</p>',
            content,
            re.DOTALL
        )
        error_match = re.search(
            r'<h3 style="[^"]*font-size: 48px[^"]*">(\d+)</h3>\s*<p style="[^"]*">Errors</p>',
            content,
            re.DOTALL
        )

        # Calculate new values
        total = int(total_match.group(1)) + total_delta if total_match else total_delta
        completed = int(completed_match.group(1)) + completed_delta if completed_match else completed_delta
        pending = int(pending_match.group(1)) + pending_delta if pending_match else pending_delta
        errors = int(error_match.group(1)) + error_delta if error_match else error_delta

        # Update statistics with new values
        # Pattern to match the number inside h3 tag before "Total Tasks"
        content = re.sub(
            r'(<h3 style="[^"]*font-size: 48px[^"]*">)\d+(</h3>\s*<p style="[^"]*">Total Tasks</p>)',
            rf'\g<1>{total}\g<2>',
            content,
            flags=re.DOTALL
        )

        # Pattern to match the number inside h3 tag before "Completed"
        content = re.sub(
            r'(<h3 style="[^"]*font-size: 48px[^"]*">)\d+(</h3>\s*<p style="[^"]*">Completed</p>)',
            rf'\g<1>{completed}\g<2>',
            content,
            flags=re.DOTALL
        )

        # Pattern to match the number inside h3 tag before "Pending"
        content = re.sub(
            r'(<h3 style="[^"]*font-size: 48px[^"]*">)\d+(</h3>\s*<p style="[^"]*">Pending</p>)',
            rf'\g<1>{pending}\g<2>',
            content,
            flags=re.DOTALL
        )

        # Pattern to match the number inside h3 tag before "Errors"
        content = re.sub(
            r'(<h3 style="[^"]*font-size: 48px[^"]*">)\d+(</h3>\s*<p style="[^"]*">Errors</p>)',
            rf'\g<1>{errors}\g<2>',
            content,
            flags=re.DOTALL
        )

        return content

    def process_single_file(self):
        """Process one file through the complete workflow."""
        # STEP 1: MONITOR
        file_path = self.monitor()
        if not file_path:
            return False

        filename = file_path.name

        # STEP 2: PROCESS
        response = self.process(file_path)
        if not response:
            self._update_dashboard(file_path.stem, "Error", filename)
            return False

        task_type = self._classify_task(self.read_file_content(file_path))

        # STEP 3: UPDATE
        if not self.update(file_path, response, task_type):
            self._update_dashboard(file_path.stem, "Error", filename)
            return False

        # STEP 4: FINALIZE
        if not self.finalize(file_path):
            return False

        return True

    def run_continuous(self, interval=5):
        """Run the agent continuously, checking for new files."""
        print(f"\n{Back.GREEN}{Fore.BLACK}{'=' * 60}{Style.RESET_ALL}")
        print(f"{Back.GREEN}{Fore.BLACK}  Agent Factory Bronze Tier - File Triage Agent  {Style.RESET_ALL}")
        print(f"{Back.GREEN}{Fore.BLACK}{'=' * 60}{Style.RESET_ALL}")
        print(f"{Fore.CYAN}Monitoring: {Fore.WHITE}{self.needs_action_path}{Style.RESET_ALL}")
        print(f"{Fore.CYAN}Supported files: {Fore.WHITE}.md, .txt, .docx{Style.RESET_ALL}")
        print(f"{Fore.CYAN}AI Provider: {Fore.WHITE}Groq API ({self.groq_model}){Style.RESET_ALL}")
        print(f"{Fore.CYAN}Check interval: {Fore.WHITE}{interval} seconds{Style.RESET_ALL}")
        print(f"{Fore.CYAN}Press Ctrl+C to stop{Style.RESET_ALL}")
        print()

        try:
            while True:
                self.process_single_file()
                time.sleep(interval)

        except KeyboardInterrupt:
            print(f"\n{Fore.YELLOW}[STOP] Agent stopped by user.{Style.RESET_ALL}")


def main():
    """Main entry point."""
    import argparse

    parser = argparse.ArgumentParser(
        description="Agent Factory Bronze Tier - File Triage Agent"
    )
    parser.add_argument(
        '--once',
        action='store_true',
        help='Process one file and exit (default: continuous mode)'
    )
    parser.add_argument(
        '--interval',
        type=int,
        default=5,
        help='Check interval in seconds for continuous mode (default: 5)'
    )

    args = parser.parse_args()

    # Initialize agent
    agent = FileTriageAgent()

    # Run in selected mode
    if args.once:
        print(f"{Fore.CYAN}[MODE] Running in single-file mode...{Style.RESET_ALL}")
        success = agent.process_single_file()
        exit(0 if success else 1)
    else:
        agent.run_continuous(interval=args.interval)


if __name__ == "__main__":
    main()
